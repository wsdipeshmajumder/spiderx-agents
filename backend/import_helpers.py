"""Knowledge ingestion helpers — pull facts from a URL (via Firecrawl) or
an uploaded file, condense to structured YAML with the best model, and fold
the result into an agent's system-prompt-backed knowledge base.

WHY YAML: the operator reviews + edits the result before it's applied, and
the agent's prompt embeds it verbatim under a clearly-bounded "KNOWLEDGE"
block — both surfaces benefit from a human-readable, hand-editable shape
that's also structured enough for the model to parse mid-call.

External deps:
  • Firecrawl HTTPS API (FIRECRAWL_API_KEY env). Used for `/v1/scrape`.
  • google-genai client (gemini-2.5-pro → gemini-2.5-flash fallback) via
    chat_bridge._best_generate, for the markdown→YAML condense pass.
  • python-docx for .docx parsing. Plain .txt is read directly. Legacy
    binary .doc, PDFs, and XLS land in PHASE 2 (we return a clear error).
"""
from __future__ import annotations

import datetime as _dt
import io as _io
import json
import logging
import os
from typing import Any, Optional

import httpx

log = logging.getLogger("eva.import")

FIRECRAWL_ENDPOINT = "https://api.firecrawl.dev/v1/scrape"
FIRECRAWL_TIMEOUT = 35.0   # site renders + Firecrawl render budget


# ─── Firecrawl scrape ──────────────────────────────────────────────────────


def _firecrawl_key() -> Optional[str]:
    key = (os.environ.get("FIRECRAWL_API_KEY") or "").strip()
    if not key or key.startswith("__"):
        return None
    return key


class IngestError(Exception):
    """Surface-able errors (HTTP-friendly messages) for ingestion routes."""

    def __init__(self, message: str, *, code: str = "ingest_failed", status: int = 502):
        super().__init__(message)
        self.code = code
        self.status = status


async def firecrawl_scrape(url: str) -> dict[str, Any]:
    """Scrape one URL via Firecrawl, returning {markdown, title, source_url}.

    Accepts any URL Firecrawl handles — websites, Google-Maps listings,
    local-listing pages. Raises IngestError with a clear message on failure
    (missing key, bad URL, Firecrawl quota, render timeout) so the caller can
    pass the reason to the operator."""
    if not isinstance(url, str) or not url.strip():
        raise IngestError("Please paste a URL.", code="bad_url", status=400)
    url = url.strip()
    if not (url.startswith("http://") or url.startswith("https://")):
        url = "https://" + url
    key = _firecrawl_key()
    if not key:
        raise IngestError(
            "FIRECRAWL_API_KEY isn't set on the server yet — add it to .env "
            "and restart to enable URL imports.",
            code="firecrawl_not_configured", status=503,
        )
    payload = {
        "url": url,
        "formats": ["markdown"],
        "onlyMainContent": True,
    }
    headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
    try:
        async with httpx.AsyncClient(timeout=FIRECRAWL_TIMEOUT) as client:
            r = await client.post(FIRECRAWL_ENDPOINT, json=payload, headers=headers)
    except httpx.TimeoutException as e:
        raise IngestError(f"That site took too long to load ({int(FIRECRAWL_TIMEOUT)}s). Try a different page.",
                          code="firecrawl_timeout", status=504) from e
    except Exception as e:  # noqa: BLE001
        raise IngestError(f"Couldn't reach Firecrawl: {e}", code="firecrawl_unreachable", status=502) from e
    if r.status_code in (401, 403):
        raise IngestError("Firecrawl rejected the API key — check FIRECRAWL_API_KEY in .env.",
                          code="firecrawl_auth", status=502)
    if r.status_code == 402:
        raise IngestError("Firecrawl says the quota is exhausted on this key. Top it up at firecrawl.dev.",
                          code="firecrawl_quota", status=502)
    if r.status_code >= 400:
        try:
            err_body = r.json()
            msg = err_body.get("error") or err_body.get("message") or r.text[:200]
        except Exception:  # noqa: BLE001
            msg = r.text[:200]
        raise IngestError(f"Firecrawl couldn't scrape that ({r.status_code}): {msg}",
                          code="firecrawl_error", status=502)
    try:
        body = r.json()
    except Exception as e:  # noqa: BLE001
        raise IngestError("Firecrawl returned a malformed response.", code="firecrawl_parse", status=502) from e
    if not body.get("success", True) or not isinstance(body.get("data"), dict):
        raise IngestError("Firecrawl couldn't extract anything useful from that URL.",
                          code="firecrawl_empty", status=422)
    data = body["data"]
    markdown = str(data.get("markdown") or "").strip()
    meta = data.get("metadata") if isinstance(data.get("metadata"), dict) else {}
    title = str(meta.get("title") or meta.get("ogTitle") or "").strip()
    if not markdown:
        raise IngestError("That page rendered but had no readable content. Try a different page.",
                          code="firecrawl_empty", status=422)
    log.info("firecrawl_scrape ok: url=%s title=%r chars=%d", url, title[:60], len(markdown))
    return {"markdown": markdown[:60_000], "title": title, "source_url": url}


# ─── LLM condense markdown → structured YAML ───────────────────────────────


_YAML_SCHEMA_HINT = """
business_name: ""        # exact spelling as the operator's listing
about: |
  Short summary of what the business does (1-2 lines).
address: ""
city: ""
country: ""
hours: ""                # in plain text, e.g. "Mon-Sat 10 AM - 8 PM, closed Sun"
phone: ""
email: ""
website: ""
services:                # list — name, description, optional price as string
  - name: ""
    description: ""
    price: ""
pricing_notes: ""
policies:                # list of strings — cancellation, deposit, age limit, etc.
  - ""
faqs:                    # list of {q, a}
  - q: ""
    a: ""
amenities:               # list of strings — parking, wifi, AC, wheelchair, etc.
  - ""
sources:                 # echo back any source URLs the page cited
  - ""
""".strip()


async def condense_to_yaml(
    *, markdown: str, source_url: str, source_title: str = "",
    context_hint: str = "", locale: str = "en-IN",
) -> str:
    """Use the best model to compress scraped/uploaded text into a structured,
    operator-reviewable YAML block. Returns the YAML text (a string).

    Includes everything the page genuinely supports; omits fields it doesn't
    mention rather than inventing them. Adds extra keys outside the schema
    when the page covers things we didn't anticipate (the catch-all spirit)."""
    # Local import to avoid a circular dep at module load.
    from . import chat_bridge

    md = (markdown or "").strip()
    if not md:
        return ""

    system = (
        "You compress one web page (or a document) into a tight, structured "
        "YAML knowledge brief a phone-AI agent will answer callers from. Rules:\n"
        "• Output VALID YAML and NOTHING else — no commentary, no code fences.\n"
        "• Use the SUGGESTED SCHEMA below where the source covers a field; "
        "OMIT fields the source doesn't mention (don't invent prices, hours, etc.).\n"
        "• You MAY add extra keys outside the schema if the page covers things "
        "the schema misses — pick concise snake_case keys for them.\n"
        "• Keep prose short. Lists are preferred over paragraphs. NEVER include "
        "marketing fluff, navigation chrome, cookie banners, or social links.\n"
        "• If the page is a Google-Maps listing, capture address, hours, "
        "phone, website, ratings summary (one line) — skip 'people also viewed'.\n"
        "SUGGESTED SCHEMA (illustrative — adapt to what the source actually has):\n"
        + _YAML_SCHEMA_HINT
    )
    parts: list[str] = []
    if context_hint:
        parts.append(f"BUSINESS CONTEXT: {context_hint.strip()}")
    parts.append(f"SOURCE URL: {source_url}")
    if source_title:
        parts.append(f"SOURCE TITLE: {source_title}")
    parts.append(f"LOCALE: {locale}")
    parts.append("CONTENT (markdown):")
    parts.append(md[:55_000])
    parts.append("\nReturn the YAML now:")
    prompt = "\n".join(parts)

    try:
        from google.genai import types as _gtypes
        client = chat_bridge._gb._client()
    except Exception as e:  # noqa: BLE001
        log.warning("condense_to_yaml: client init failed: %s", e)
        return ""

    cfg = _gtypes.GenerateContentConfig(
        system_instruction=system,
        temperature=0.2,
        # YAML — don't force application/json here, just plain text.
    )
    models = [chat_bridge.CATCHALL_MODEL]
    if chat_bridge.CHAT_MODEL != chat_bridge.CATCHALL_MODEL:
        models.append(chat_bridge.CHAT_MODEL)
    import asyncio as _asyncio
    text: str = ""
    for model in models:
        try:
            resp = await _asyncio.wait_for(
                client.aio.models.generate_content(model=model, contents=prompt, config=cfg),
                timeout=40.0,
            )
            t = getattr(resp, "text", None) or ""
            t = t.strip()
            # Strip code fences if the model added them anyway.
            if t.startswith("```"):
                t = t.split("\n", 1)[-1] if "\n" in t else t
                if t.endswith("```"):
                    t = t[: t.rfind("```")].strip()
                if t.lower().startswith("yaml"):
                    t = t.split("\n", 1)[-1] if "\n" in t else t
            if t:
                text = t.strip()
                log.info("condense_to_yaml: %d bytes (model=%s, source=%s)", len(text), model, source_url[:80])
                break
        except Exception as e:  # noqa: BLE001
            log.warning("condense_to_yaml(%s) failed: %s", model, e)
            continue
    return text


# ─── File parsing (knowledge-base upload) ──────────────────────────────────


_SUPPORTED_TEXT_EXT = {".txt", ".md", ".markdown"}
_SUPPORTED_DOCX_EXT = {".docx"}
_COMING_SOON_EXT = {".pdf", ".doc", ".xls", ".xlsx", ".csv"}


def parse_file_to_text(filename: str, raw: bytes) -> str:
    """Decode an uploaded file into plain text we can hand to the YAML
    condenser. Raises IngestError with a friendly reason on failure."""
    name = (filename or "").lower().strip()
    if not name:
        raise IngestError("File has no name.", code="bad_file", status=400)
    if not raw:
        raise IngestError("That file was empty.", code="empty_file", status=400)
    if len(raw) > 8 * 1024 * 1024:
        raise IngestError("File is over 8 MB — split it into smaller pieces.", code="too_large", status=413)

    ext = ""
    if "." in name:
        ext = "." + name.rsplit(".", 1)[-1]

    if ext in _SUPPORTED_TEXT_EXT:
        # Best-effort decode: utf-8 first, then latin-1 as a fallback.
        for enc in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return raw.decode(enc).strip()
            except UnicodeDecodeError:
                continue
        raise IngestError("Couldn't read that text file — please re-save as UTF-8.",
                          code="decode_error", status=422)

    if ext in _SUPPORTED_DOCX_EXT:
        try:
            import docx  # python-docx
        except Exception as e:  # noqa: BLE001
            raise IngestError("Word-doc parser isn't installed on the server (python-docx).",
                              code="docx_missing", status=503) from e
        try:
            f = _io.BytesIO(raw)
            doc = docx.Document(f)
            lines: list[str] = []
            for p in doc.paragraphs:
                t = (p.text or "").strip()
                if t:
                    lines.append(t)
            # Tables (menus, price lists) often live in tables in Word docs.
            for tbl in doc.tables:
                for row in tbl.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    cells = [c for c in cells if c]
                    if cells:
                        lines.append(" | ".join(cells))
            text = "\n".join(lines).strip()
            if not text:
                raise IngestError("That .docx had no readable text.", code="docx_empty", status=422)
            return text
        except IngestError:
            raise
        except Exception as e:  # noqa: BLE001
            raise IngestError(f"Couldn't read that .docx: {e}", code="docx_parse_error", status=422) from e

    if ext in _COMING_SOON_EXT:
        nice = ext.lstrip(".").upper()
        raise IngestError(f"{nice} files aren't supported yet — coming soon. Try .txt or .docx for now.",
                          code="format_coming_soon", status=415)

    raise IngestError(
        f"Unsupported file type {ext or '(unknown)'}. Supported now: .txt, .md, .docx.",
        code="unsupported_format", status=415,
    )


# ─── Folding YAML into an agent's knowledge ────────────────────────────────


_KNOWLEDGE_HEADER = "━━━━━━━━━━━━━ KNOWLEDGE PULLED FROM SOURCES ━━━━━━━━━━━━━"
_KNOWLEDGE_FOOTER = "━━━━━━━━━━━━━ END KNOWLEDGE FROM SOURCES ━━━━━━━━━━━━━"


def append_knowledge_block(existing_prompt: str, yaml_text: str, source: dict[str, Any]) -> str:
    """Append (or rebuild) the KNOWLEDGE block on an agent's system prompt.

    Multiple imports stack INSIDE one block so the prompt has a single, clear
    'reference facts' region. Returns the full new system prompt string."""
    yaml_text = (yaml_text or "").strip()
    if not yaml_text:
        return existing_prompt or ""
    label = source.get("title") or source.get("source") or source.get("url") or "Source"
    when = _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%d")
    block_entry = f"### {label} ({when})\n```yaml\n{yaml_text}\n```"
    body = (existing_prompt or "").rstrip()
    if _KNOWLEDGE_HEADER in body and _KNOWLEDGE_FOOTER in body:
        # Insert before the footer so the block grows in order of arrival.
        head, rest = body.split(_KNOWLEDGE_FOOTER, 1)
        new_body = head.rstrip() + "\n\n" + block_entry + "\n\n" + _KNOWLEDGE_FOOTER + rest
        return new_body
    # First import — open a fresh block.
    intro = (
        "These reference facts come from sources the operator uploaded or "
        "imported. ANSWER CALLERS USING THESE — do not invent prices, hours, "
        "or availability beyond what's stated here. If a caller asks about "
        "something not covered, offer to take a callback rather than guess."
    )
    block = "\n\n".join([_KNOWLEDGE_HEADER, intro, block_entry, _KNOWLEDGE_FOOTER])
    return (body + ("\n\n" if body else "") + block).rstrip()


def add_source_to_variables(
    variables: dict[str, Any] | None, source: dict[str, Any],
) -> dict[str, Any]:
    """Append a knowledge-source record to variables.knowledge_sources so the
    dashboard can list what was imported and when."""
    v = dict(variables or {})
    items = v.get("knowledge_sources")
    if not isinstance(items, list):
        items = []
    rec = {
        "kind": source.get("kind") or "url",
        "url": source.get("url") or "",
        "title": source.get("title") or "",
        "filename": source.get("filename") or "",
        "added_at": _dt.datetime.now(_dt.timezone.utc).isoformat(timespec="seconds"),
    }
    items.append(rec)
    v["knowledge_sources"] = items[-50:]   # cap so it never balloons
    return v
